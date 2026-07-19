from __future__ import annotations

import argparse
import csv
from datetime import date
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
FIGURES = REPORTS / "figures"
OUTPUT = REPORTS / "导师汇报_A股多因子选股策略研究.docx"
RUN_DIR: Path | None = None
RUN_METADATA: dict[str, object] = {}

# Resolved preset: standard_business_brief.
# Named overrides: Microsoft YaHei for East-Asian glyphs and an editorial-cover
# title block suitable for a formal academic progress report.
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
MUTED = "666666"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "F4F7FA"
RISK_RED = "9B1C1C"
CAUTION = "7A5A00"
WHITE = "FFFFFF"
BLACK = "000000"
BODY_FONT = "Calibri"
CJK_FONT = "Microsoft YaHei"
CONTENT_WIDTH_DXA = 9026
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    original_total = sum(widths_dxa)
    widths_dxa = [round(width * CONTENT_WIDTH_DXA / original_total) for width in widths_dxa]
    widths_dxa[-1] += CONTENT_WIDTH_DXA - sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size: float | None = None, bold: bool | None = None,
                 color: str | None = None, italic: bool | None = None) -> None:
    run.font.name = BODY_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), BODY_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), BODY_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph_runs(paragraph, size: float = 11, color: str = BLACK) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def set_paragraph_shading(paragraph, fill: str, border_color: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border_color:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "6")
        left.set(qn("w:color"), border_color)
        p_bdr.append(left)


def set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_body(doc: Document, text: str, *, bold_lead: str | None = None,
             color: str = BLACK, after: float = 6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=11, bold=True, color=color)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=11, color=color)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=color)
    return p


def add_callout(doc: Document, label: str, text: str, *, fill: str = PALE_BLUE,
                accent: str = BLUE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.10
    set_paragraph_shading(p, fill, accent)
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, size=11, bold=True, color=accent)
    r2 = p.add_run(text)
    set_run_font(r2, size=11, color=BLACK)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(p)
    style_paragraph_runs(p, size={1: 16, 2: 13, 3: 12}[level],
                         color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level])


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=9, color=MUTED, italic=True)


def add_picture(doc: Document, filename: str, caption: str, width: float = 6.2) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(FIGURES / filename), width=Inches(width))
    add_caption(doc, caption)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run("第 ")
    set_run_font(r, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    r2 = paragraph.add_run(" 页")
    set_run_font(r2, size=9, color=MUTED)


def load_metrics() -> dict[str, float]:
    rows: dict[str, float] = {}
    source = RUN_DIR / "metrics.csv" if RUN_DIR is not None else FIGURES / "performance_metrics.csv"
    with source.open(encoding="utf-8-sig", newline="") as fh:
        for row in csv.DictReader(fh):
            key = row.get("") or row.get("metric") or next(iter(row.values()))
            rows[str(key)] = float(row["value"])
    return rows


def load_first_row(filename: str) -> dict[str, str]:
    path = FIGURES / filename
    if not path.exists():
        return {}
    with path.open(encoding="utf-8-sig", newline="") as fh:
        return next(csv.DictReader(fh), {})


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int],
              *, header_fill: str = LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=9.5, bold=True, color=NAVY)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            set_run_font(r, size=9.5, color=BLACK)
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(4)
    return table


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for level, size, before, after, color in (
        (1, 16, 16, 8, BLUE),
        (2, 13, 12, 6, BLUE),
        (3, 12, 8, 4, DARK_BLUE),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(96)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("阶段性研究报告")
    set_run_font(r, size=12, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("A股多因子选股策略研究")
    set_run_font(r, size=28, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(56)
    r = p.add_run("可审计研究框架的设计、实现与阶段性验证")
    set_run_font(r, size=15, color=DARK_BLUE)

    for label, value in (
        ("导师", "李老师"),
        ("汇报人", "Jason Chen"),
        ("项目版本", f"main @ {str(RUN_METADATA.get('git_commit', '58b7521'))[:7]}"),
        ("证据基线", str(RUN_METADATA.get("run_id", "advisor-baseline-20260713"))),
        ("报告日期", date.today().isoformat()),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        r1 = p.add_run(f"{label}：")
        set_run_font(r1, size=11, bold=True, color=MUTED)
        r2 = p.add_run(value)
        set_run_font(r2, size=11, color=BLACK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(54)
    p.paragraph_format.space_after = Pt(0)
    set_paragraph_shading(p, LIGHT_GRAY)
    r = p.add_run("研究性质：工程与方法验证。当前结果基于合成样例数据，不构成真实市场有效性或可交易性证明。")
    set_run_font(r, size=10, bold=True, color=RISK_RED)


def configure_body_section(doc: Document):
    body = doc.add_section(WD_SECTION.NEW_PAGE)
    body.page_width = Mm(210)
    body.page_height = Mm(297)
    body.top_margin = Inches(1)
    body.bottom_margin = Inches(1)
    body.left_margin = Inches(1)
    body.right_margin = Inches(1)
    body.header_distance = Inches(0.492)
    body.footer_distance = Inches(0.492)
    body.header.is_linked_to_previous = False
    body.footer.is_linked_to_previous = False
    hp = body.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("A股多因子选股策略研究｜阶段性研究报告")
    set_run_font(r, size=9, color=MUTED)
    fp = body.footer.paragraphs[0]
    add_page_field(fp)
    return body


def configure_run_source(run_dir: str | Path | None) -> None:
    global RUN_DIR, FIGURES, RUN_METADATA
    if run_dir is None:
        return
    RUN_DIR = Path(run_dir).resolve()
    if not RUN_DIR.exists():
        raise FileNotFoundError(RUN_DIR)
    FIGURES = RUN_DIR / "figures"
    metadata_path = RUN_DIR / "run_metadata.json"
    if not metadata_path.exists():
        raise FileNotFoundError(metadata_path)
    RUN_METADATA = json.loads(metadata_path.read_text(encoding="utf-8"))


def build(run_dir: str | Path | None = None, output_path: str | Path | None = None) -> Path:
    configure_run_source(run_dir)
    metrics = load_metrics()
    cost_row = load_first_row("cost_attribution.csv")
    unfilled_row = load_first_row("unfilled_order_analysis.csv")
    unfilled_count = int(float(unfilled_row.get("order_count", 0)))
    unfilled_value = float(unfilled_row.get("unfilled_value", 0.0))
    total_cost_amount = float(cost_row.get("total_cost", 0.0))
    total_cost_ratio = float(cost_row.get("total_cost_ratio", 0.0))
    reconciliation_residual = float(cost_row.get("reconciliation_residual", 0.0))
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    configure_body_section(doc)

    add_heading(doc, "摘要", 1)
    add_body(
        doc,
        "本项目面向A股多因子选股研究中常见的时间错配、数据泄漏、样本外污染和交易不可达问题，"
        "构建了一套从标准化数据接入、Point-in-Time（PIT）因子生成、滚动样本外验证，到受约束组合构建、"
        "事件驱动回测、绩效归因与LLM事件标签审计的可复现工程框架。当前版本已打通数据、因子、组合、"
        "执行、归因和报告链路，并通过完整质量门禁。"
    )
    add_body(
        doc,
        "阶段性验证全部使用项目内合成样例数据。标准pipeline的样例净收益为"
        f"{metrics['net_total_return']:.2%}，最大回撤为{metrics['max_drawdown']:.2%}，Sharpe为"
        f"{metrics['sharpe']:.3f}；这些数值仅证明指标计算与审计产物能够稳定生成，不能用于判断因子在真实市场中的有效性。"
        "当前最主要的研究缺口是尚未在完整历史PIT数据上形成有效的滚动样本外结果，同时部分长周期因子覆盖不足。"
        "成本归因已经统一为货币金额和期初净值比例两套字段，并保留与净值成本拖累的勾稽差。"
    )
    add_callout(
        doc,
        "核心结论",
        "项目已经完成“可信研究基础设施”的第一阶段建设；下一阶段的决定性任务不是继续扩展因子数量，"
        "而是接入可核验的真实PIT数据，并以样本外、执行延迟、成本和容量压力测试建立证据。"
    )

    add_heading(doc, "一、研究背景与问题定义", 1)
    add_body(
        doc,
        "多因子选股的难点不只在于构造信号，还在于确认信号所用信息在当时是否可见、组合能否按假设成交、"
        "收益是否来自少数行业或市值暴露，以及结论是否经得住样本外和成本扰动。若忽略这些环节，"
        "回测结果即使表现优异，也可能由未来函数、幸存者偏差、基准错配或费用缺失造成。"
    )
    add_body(
        doc,
        "本研究因此将问题定义为：在明确A股交易规则与数据时点的前提下，建立一套能够复现、审计和扩展的多因子研究流程，"
        "并将“因子统计有效性”“组合回测表现”和“真实可执行性”分层评价。当前研究对象为中证500风格股票池，"
        "配置基准为000905.SH，默认月末调仓，因子在信号日收盘后生成，下一交易日开盘尝试执行。"
    )

    add_heading(doc, "二、研究设计与技术路线", 1)
    add_heading(doc, "2.1 时间线与防泄漏约束", 2)
    add_table(
        doc,
        ["环节", "当前定义", "审计重点"],
        [
            ["信号生成", "T日收盘后", "仅使用T日及之前可见数据"],
            ["财务数据", "usable_date到达后", "保留报告期、公告日、可用日"],
            ["交易执行", "T+1交易日开盘", "停牌、涨跌停、手数和参与率约束"],
            ["目标收益", "未来20个交易日", "仅用于评价，不回填信号"],
            ["样本外验证", "训练/验证/测试滚动", "方向和权重在测试窗口前锁定"],
        ],
        [1800, 3000, 4560],
    )
    add_body(
        doc,
        "该时间线将数据时间、信号时间、订单时间、成交时间和收益归因窗口拆开。全样本IC只作为诊断展示，"
        "不用于同期组合权重；真实模式在缺少数据manifest、关键PIT表或可用样本外评分时执行阻断。"
    )

    add_heading(doc, "2.2 因子体系与截面处理", 2)
    add_body(
        doc,
        "当前因子覆盖量价、风险、估值、质量、成长、资金流和事件信息，共31个配置项。处理流程在每个交易日横截面内依次执行"
        "MAD去极值、z-score标准化，并按配置进行行业和市值中性化。因子定义、输入字段、方向、窗口与PIT要求均形成结构化规格表。"
    )
    add_table(
        doc,
        ["类别", "代表因子", "主要风险"],
        [
            ["量价/风险", "动量、反转、波动率、Beta、流动性", "窗口不足、停牌与复权口径"],
            ["估值/基本面", "BP、EP、SP、CFP、规模", "负值处理、公告时点和修订值"],
            ["质量/成长", "ROE、ROA、毛利率、收入与利润增长", "财务可用日和行业暴露"],
            ["资金流", "5/20日资金流与大单资金流", "数据源稳定性和零成交"],
            ["事件", "20日事件情绪", "发布时间、去重、标签一致性"],
        ],
        [1800, 3000, 4560],
    )

    add_heading(doc, "2.3 组合、执行与归因", 2)
    add_body(
        doc,
        "默认组合目标为Top50，单票上限5%，最少持仓30只，单行业上限30%，最大现金10%，单次换手上限50%。"
        "事件驱动回测保留订单、成交、持仓和未成交原因，并拆分佣金、印花税、滑点与冲击成本。"
        "绩效模块同时计算绝对收益、相对基准表现、回撤、行业/市值/个股贡献和成本拖累。"
    )

    add_heading(doc, "三、阶段性实现成果", 1)
    add_table(
        doc,
        ["模块", "已实现能力", "可复核产物"],
        [
            ["数据工程", "12张标准表、导入映射、主键与跨表质量检查", "manifest与质量报告"],
            ["PIT控制", "信号/执行/目标日期分离，财务可用日追踪", "timing审计表与测试"],
            ["因子研究", "覆盖率、IC、分组、衰减、相关性与中性化", "CSV与同源图表"],
            ["组合回测", "约束组合、T+1执行、订单/成交/持仓审计", "orders、fills、positions"],
            ["稳健性", "成本、延迟与成交额参与率情景", "scenario与未成交分析"],
            ["绩效归因", "绝对/相对指标、行业/市值/回撤/成本归因", "metrics与归因表"],
            ["LLM事件", "离线标签、版本、缓存、抽查门槛", "标签与审计报告"],
            ["工程质量", "编译、单测、CLI和Notebook smoke", "统一质量门禁"],
        ],
        [1600, 4460, 3300],
    )
    add_callout(
        doc,
        "质量验证",
        "2026-07-13在main@58b7521及本轮源码树快照上完成复核：57个单元测试通过；7个Notebook smoke通过；"
        "CLI样例生成和标准pipeline smoke通过。该结论仅针对代码与合成样例链路。",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "四、样例结果与审慎解释", 1)
    add_heading(doc, "4.1 标准pipeline绩效", 2)
    add_table(
        doc,
        ["指标", "样例结果", "解释边界"],
        [
            ["累计净收益", f"{metrics['net_total_return']:.2%}", "合成样例，仅验证计算链路"],
            ["年化收益", f"{metrics['annual_return']:.2%}", "180个交易日样本，年化敏感"],
            ["年化波动率", f"{metrics['annual_volatility']:.2%}", "基于252交易日年化"],
            ["Sharpe", f"{metrics['sharpe']:.3f}", "不支持策略有效性结论"],
            ["最大回撤", f"{metrics['max_drawdown']:.2%}", "样例路径风险展示"],
            ["成本拖累", f"{metrics['cost_drag']:.2%}", "标准成本情景，非真实成交校准"],
            ["全期平均实际持仓数", f"{metrics['avg_holding_count']:.2f}", f"含{int(metrics['pre_investment_days'])}个未建仓日，不用于判断约束失效"],
            ["建仓后平均实际持仓数", f"{metrics['avg_holding_count_invested_days']:.2f}", f"目标均值{metrics['avg_target_holding_count']:.0f}只；成交后持仓约束通过"],
            ["建仓后平均现金比例", f"{metrics['avg_cash_weight_invested_days']:.2%}", f"最大{metrics['max_cash_weight_invested_days']:.2%}，低于10%上限"],
        ],
        [2100, 1800, 5460],
    )
    add_picture(doc, "cumulative_return.png", "图1  标准pipeline合成样例累计净值。前段持平来自组合尚未形成或未执行，不应解释为低风险。")
    add_picture(doc, "drawdown.png", "图2  标准pipeline合成样例回撤。该图用于验证回撤计算与区间识别，不代表真实市场风险水平。")

    add_heading(doc, "4.2 因子诊断", 2)
    add_body(
        doc,
        "合成样例中部分因子呈现较高Rank IC，并已同时输出普通t值、以19阶Newey–West标准误计算的HAC t值及"
        "Benjamini–Hochberg FDR。由于样例数据由规则生成，这些结果可能反映生成机制本身，不能外推到真实A股。"
    )
    add_body(
        doc,
        "每日20日前瞻收益存在重叠，现只保留为横截面诊断，不再直接连乘生成累计曲线。月末非重叠分组检验"
        "已单独输出group_test_nonoverlap.csv；本样例仅有2个合法持有期，样本量不足，不能据此作市场有效性结论。"
    )
    add_body(
        doc,
        "覆盖率审计显示，mom_60_skip5与mom_120在当前因子面板中无有效覆盖，idio_vol_60覆盖率仅约3.3%。"
        "因此长窗口因子尚不具备有效比较基础。当前walk-forward输出中的样本外IC与权重表为空，说明样例历史长度不足以形成正式滚动样本外证据。"
    )
    add_picture(doc, "ic_series.png", "图3  合成样例综合score的Rank IC序列（n=40）。序列波动与尾段转弱提示不能只看均值。")

    add_heading(doc, "4.3 执行与成本审计", 2)
    add_body(
        doc,
        f"标准pipeline记录了{unfilled_count}笔因手数约束未成交的订单，请求金额合计约{unfilled_value / 10_000:.2f}万元。"
        "全期平均实际持仓14.21只的原因是"
        f"180个观测日中有{int(metrics['pre_investment_days'])}个未建仓日；建仓后平均实际持仓为"
        f"{metrics['avg_holding_count_invested_days']:.2f}只，目标组合均值为{metrics['avg_target_holding_count']:.0f}只。"
        "成交后单票、行业、现金、换手和参与率审计在该样例中均无违例，因此不能再用全期均值认定配置失效。"
    )
    add_body(
        doc,
        "cost_attribution.csv现已分别给出commission_amount、stamp_tax_amount、slippage_amount、impact_amount等货币字段，"
        f"以及对应的期初净值比例字段。样例总成本为{total_cost_amount:,.2f}元，占期初资产{total_cost_ratio:.4%}；"
        f"净值成本拖累为{metrics['cost_drag']:.4%}，两者因成本发生时点和复利形成约"
        f"{reconciliation_residual:.4%}勾稽差。"
    )

    add_heading(doc, "五、可信度审计与研究边界", 1)
    add_table(
        doc,
        ["结论或主张", "当前证据", "判定"],
        [
            ["工程链路可复现", "单测、Notebook、CLI与pipeline门禁均通过", "已支持"],
            ["时间语义已显式化", "信号/执行/目标日期与财务可用日可审计", "已支持"],
            ["因子在真实A股有效", "仅有合成样例IC和分组结果", "未支持"],
            ["策略具有样本外稳定性", "正式OOS IC与权重输出为空", "未支持"],
            ["组合可按假设成交", "有订单审计，但缺真实盘口与冲击校准", "部分支持"],
            ["LLM事件因子可入组合", "reviewed_count=0，尚无人审通过率", "未支持"],
            ["绩效归因口径一致", "成本金额、期初资产比例与净值拖累均保留并可勾稽", "已支持"],
        ],
        [3000, 4560, 1800],
    )
    add_callout(
        doc,
        "最重要的剩余风险",
        "历史指数成分、退市、新股、ST、停复牌、涨跌停、财务修订和事件文本尚未由可靠的真实PIT数据覆盖。"
        "在这些数据通过质量阻断和时间一致性审计之前，不应把样例结果表述为投资结论。",
        fill="FFF6E5",
        accent=CAUTION,
    )

    add_heading(doc, "六、下一阶段研究计划", 1)
    add_table(
        doc,
        ["阶段", "工作重点", "验收标准"],
        [
            ["阶段A：真实数据闭环", "接入中证500历史成分、复权行情、财务可用日与交易状态", "manifest完整；阻断问题为0；覆盖率按年稳定"],
            ["阶段B：样本外证据", "形成滚动训练/验证/测试，锁定方向和权重", "OOS IC非空；分年度/市场状态稳定；无跨窗泄漏"],
            ["阶段C：执行压力测试", "成本、延迟、参与率、资金规模与未成交再平衡", "净收益对成本和延迟不过度敏感；成交后约束合格"],
            ["阶段D：论文式总结", "完成基准、消融、稳健性和限制讨论", "每项主张有对应表图与可复现run目录"],
        ],
        [2100, 3900, 3360],
    )
    add_body(
        doc,
        "建议优先完成阶段A和阶段B，再决定是否扩展因子库或引入复杂优化器。若在真实PIT数据上不能得到稳定的样本外Rank IC、"
        "分组单调性和可承受的换手成本，应保留框架但重新评估信号，而不是通过反复调参追求更好的历史曲线。"
    )

    add_heading(doc, "七、结论", 1)
    add_body(
        doc,
        "本项目已完成一套面向A股多因子研究的可审计工程框架，并把数据版本、PIT时点、因子处理、样本外流程、"
        "组合约束、订单成交、交易成本、风险归因和LLM标签审计纳入统一链路。完整质量门禁表明代码与合成样例流程可稳定复现。"
        "但当前证据仍停留在工程和方法验证阶段，尚不足以证明真实市场有效性、样本外稳定性或可交易性。"
        "下一阶段应以真实PIT数据和严格样本外研究为主线，形成可由导师逐项复核的结论—证据—边界闭环。"
    )

    add_heading(doc, "附录：复现入口与主要产物", 1)
    add_table(
        doc,
        ["用途", "入口或产物"],
        [
            ["配置契约", "python -m ashare_factor_research.main validate-config"],
            ["冻结协议运行", "python -m ashare_factor_research.main run-research --protocol config/research_protocol.yaml"],
            ["数据验证", "python -m ashare_factor_research.main verify-data --data-dir <path> --mode real"],
            ["指定run生成报告", "python -m ashare_factor_research.main build-advisor-report --run-dir <path>"],
            ["研究Notebook", "notebooks/01_data_collection.ipynb 至 07_llm_event_analysis.ipynb"],
            ["标准报告", "reports/factor_research_report.md 与 PDF"],
            ["图表与底表", "reports/figures/"],
            ["实现状态", "reports/implementation_status.md"],
        ],
        [2600, 6760],
    )
    add_body(
        doc,
        "说明：Notebook 05使用Top10、单票上限20%的演示参数，标准pipeline使用Top50、单票上限5%。"
        "两者用于不同的smoke场景，导师审阅时应以标准pipeline及其配置快照为统一口径。",
        color=RISK_RED,
    )

    doc.core_properties.title = "A股多因子选股策略研究：可审计研究框架的设计、实现与阶段性验证"
    doc.core_properties.subject = "导师阶段性研究汇报"
    doc.core_properties.author = "Jason Chen"
    doc.core_properties.keywords = "A股, 多因子, PIT, 样本外, 回测, 绩效归因"
    target = Path(output_path).resolve() if output_path else OUTPUT
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)
    return target


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--run-dir")
    parser.add_argument("--output")
    args = parser.parse_args()
    print(build(args.run_dir, args.output))
